import email
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def get_html_from_mhtml(mhtml_path):
    """Trích xuất nội dung HTML từ tệp MHTML."""
    with open(mhtml_path, 'r', encoding='utf-8') as f:
        msg = email.message_from_file(f)
        for part in msg.walk():
            if part.get_content_type() == 'text/html':
                charset = part.get_content_charset() or 'utf-8'
                return part.get_payload(decode=True).decode(charset)
    return None

def process_element(element, doc, current_paragraph=None, style_stack=None):
    """
    Xử lý đệ quy từng phần tử HTML và thêm vào tài liệu DOCX.
    """
    if style_stack is None:
        style_stack = []

    # Bỏ qua các thẻ không cần thiết
    if element.name in ['script', 'style', 'meta', 'link']:
        return

    # Xử lý các chuỗi văn bản
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text and current_paragraph is not None:
            run = current_paragraph.add_run(text)
            # Áp dụng các style hiện tại (đậm, nghiêng)
            for style in style_stack:
                if style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
        return

    if not hasattr(element, 'name'):
        return

    # Xử lý các thẻ cụ thể
    tag_name = element.name.lower()
    
    # Xử lý công thức LaTeX (cả inline và display)
    if tag_name == 'span' and 'katex' in element.get('class', []):
        latex_annotation = element.find('annotation')
        if latex_annotation:
            latex_code = latex_annotation.get_text(strip=True)
            # Thêm công thức vào một đoạn mới cho rõ ràng
            p = doc.add_paragraph()
            run = p.add_run(f"Công thức LaTeX: {latex_code}")
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return # Dừng xử lý các thẻ con của katex

    # Thẻ khối tạo đoạn mới
    if tag_name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div']:
        text_content = element.get_text(strip=True)
        if not text_content: # Bỏ qua các thẻ div/p rỗng
             # Nhưng vẫn xử lý con của chúng để không bỏ sót nội dung lồng nhau
            for child in element.children:
                process_element(child, doc, current_paragraph, style_stack)
            return

        if tag_name.startswith('h'):
            level = int(tag_name[1])
            p = doc.add_heading(level=level)
        else:
            p = doc.add_paragraph()

        # Xử lý các phần tử con của thẻ này
        for child in element.children:
            process_element(child, doc, p, style_stack)
        return

    # Thẻ danh sách
    if tag_name in ['ul', 'ol']:
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            # Xử lý nội dung bên trong mỗi mục li
            for child in li.children:
                process_element(child, doc, p, style_stack)
        return

    # Thẻ định dạng inline
    if tag_name in ['strong', 'b']:
        style_stack.append('bold')
    elif tag_name in ['em', 'i']:
        style_stack.append('italic')

    # Đệ quy xử lý các thẻ con
    for child in element.children:
        process_element(child, doc, current_paragraph, style_stack)
    
    # Sau khi xử lý xong các con, xóa style khỏi stack
    if tag_name in ['strong', 'b', 'em', 'i']:
        style_stack.pop()

def html_to_docx(html_content, docx_path):
    """
    Hàm chính để chuyển đổi HTML sang DOCX.
    """
    if not html_content:
        print("Nội dung HTML trống.")
        return
        
    soup = BeautifulSoup(html_content, 'lxml')
    doc = Document()
    
    # Bắt đầu xử lý từ thẻ body
    if soup.body:
        process_element(soup.body, doc)
    else:
        print("Không tìm thấy thẻ <body> trong HTML.")
        return

    doc.save(docx_path)
    print(f"Tệp DOCX đã được lưu tại: {docx_path}")

if __name__ == '__main__':
    mhtml_file = 'ChatGPT_TTNT.mhtml'  # Thay thế bằng tên tệp MHTML của bạn
    docx_file = 'ChatGPT_TTNT_output.docx' # Tên tệp DOCX đầu ra mới

    html_data = get_html_from_mhtml(mhtml_file)
    
    html_to_docx(html_data, docx_file)