import email
from bs4 import BeautifulSoup
from docx import Document
import re

def get_html_from_mhtml(mhtml_path):
    """
    Trích xuất nội dung HTML từ tệp MHTML.

    Args:
        mhtml_path (str): Đường dẫn đến tệp MHTML.

    Returns:
        str: Nội dung HTML dưới dạng chuỗi.
    """
    with open(mhtml_path, 'r', encoding='utf-8') as f:
        msg = email.message_from_file(f)
        for part in msg.walk():
            if part.get_content_type() == 'text/html':
                # Sửa lỗi: Cung cấp 'utf-8' làm bảng mã mặc định nếu không tìm thấy
                charset = part.get_content_charset() or 'utf-8'
                return part.get_payload(decode=True).decode(charset)
    return None

def add_latex_to_doc(doc, latex_string):
    """
    Thêm một chuỗi LaTeX vào tài liệu Word dưới dạng văn bản.

    Args:
        doc (docx.document.Document): Đối tượng tài liệu Word.
        latex_string (str): Chuỗi LaTeX cần thêm.
    """
    # Loại bỏ các ký tự không mong muốn và chuẩn hóa chuỗi LaTeX
    clean_latex = re.sub(r'\\math(cal|bf|it|rm|sf|tt)', '', latex_string).strip()
    p = doc.add_paragraph()
    p.add_run(f'Công thức LaTeX: {clean_latex}')

def html_to_docx(html_content, docx_path):
    """
    Chuyển đổi nội dung HTML với công thức LaTeX sang tài liệu DOCX.

    Args:
        html_content (str): Nội dung HTML cần chuyển đổi.
        docx_path (str): Đường dẫn để lưu tệp DOCX.
    """
    soup = BeautifulSoup(html_content, 'lxml')
    doc = Document()

    # Tìm tất cả các thẻ có liên quan đến nội dung
    content_elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'div'])

    for element in content_elements:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            doc.add_heading(element.get_text(), level=int(element.name[1]))
        elif element.name == 'p':
            # Tìm các công thức toán học trong đoạn văn
            formulas = element.find_all('span', class_='katex-display')
            if formulas:
                for formula in formulas:
                    # Trích xuất mã LaTeX gốc từ thẻ annotation
                    latex_annotation = formula.find('annotation')
                    if latex_annotation:
                        latex_code = latex_annotation.get_text(strip=True)
                        add_latex_to_doc(doc, latex_code)
                    formula.decompose() # Xóa công thức đã xử lý

            # Thêm văn bản còn lại của đoạn văn
            text = element.get_text(strip=True)
            if text:
                doc.add_paragraph(text)

        elif element.name == 'ul' or element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(strip=True), style='List Bullet')

    doc.save(docx_path)
    print(f"Tệp DOCX đã được lưu tại: {docx_path}")

if __name__ == '__main__':
    # Đường dẫn đến tệp MHTML của bạn
    mhtml_file = 'Gemini_latex3a.mhtml' # Thay thế bằng tên tệp của bạn
    # Đường dẫn để lưu tệp DOCX đầu ra
    docx_file = 'Gemini_latex3a_output.docx'

    # Trích xuất HTML từ MHTML
    html_data = get_html_from_mhtml(mhtml_file)

    if html_data:
        # Chuyển đổi HTML sang DOCX
        html_to_docx(html_data, docx_file)
    else:
        print("Không tìm thấy nội dung HTML trong tệp MHTML.")